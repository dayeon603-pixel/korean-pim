#!/usr/bin/env python3
"""Estimate the typeset page count of the manuscript — python3 analysis/page_estimate.py

Why this exists. The submission has a five-page limit and the manuscript was trimmed four times
against a character-per-line guess that turned out to be wrong by roughly forty per cent. Guessing
cost real content. This script measures instead: it lays out each paragraph with actual font
metrics at the specification's point size, character scaling and tracking, then adds the paragraph
spacing the specification prescribes.

Limits, which matter. The submission font is 신명조, a Korean face; its Latin glyphs are not
measured here because the file is not available to this script, so Times New Roman is used as the
proxy for Latin text. A wider Latin face would raise the count. Word remains the authority and this
script is a way to stop trimming blind, not a substitute for opening the file.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from PIL import ImageFont

# F7 page geometry from the society's specification.
PAGE_HEIGHT_MM = 297.0
MARGIN_TOP_MM = MARGIN_BOTTOM_MM = 10.0
MARGIN_SIDE_MM = 25.0
PAGE_WIDTH_MM = 210.0

USABLE_HEIGHT_MM = PAGE_HEIGHT_MM - MARGIN_TOP_MM - MARGIN_BOTTOM_MM
LINE_WIDTH_MM = PAGE_WIDTH_MM - 2 * MARGIN_SIDE_MM

CHAR_SCALE = 0.95      # 장평 95%
CHAR_SPACING = -0.05   # 자간 -5%, applied per character as a fraction of the point size
MEASURE_PT = 200       # measure large, then scale down, to keep rounding out of the result

PROXY_FONT = "/System/Library/Fonts/Supplemental/Times New Roman.ttf"
import os
DOCX = Path(os.environ.get("PAPER_DOCX",
    Path.home() / "Documents" / "Yakson" / "학회제출" / "KOSMI2026_admin_condition_axis.docx"))

# Row height and per-table padding, in mm. Measured from the generated tables at 9pt.
TABLE_ROW_MM = 5.5
TABLE_PAD_MM = 6.0

# A single line box is not the point size. For Times New Roman the hhea metrics give
# (ascent - descent + lineGap) / unitsPerEm = 1.150 em, so 10pt text sets on an 11.5pt line before
# leading is applied. The first version of this script used the point size itself and understated
# every paragraph by fifteen per cent.
LINE_BOX_EM = 1.150

# Two page counts have now been measured in Word, and one scalar cannot fit both:
#
#   conference version  19,882 chars,  3 tables  ->  model 4.58,  Word 5 pages
#   journal version     28,267 chars,  6 tables  ->  model 6.56,  Word 8 pages
#
# Probing the conference file with filler text put its true fill at 4.31 to 4.45 pages, which
# implies a factor near 0.96, while the journal implies at least 1.07. The gap tracks table count,
# so the flat TABLE_PAD_MM under-costs a table and no single factor absorbs it. The model is
# therefore a fallback only. Word is the authority and is queried first; see word_pages().
CALIBRATION = 1.0


def word_pages(path: Path) -> int | None:
    """Ask Word for the real page count. Returns None if Word is unavailable or refuses the file.

    Word is sandboxed out of temporary directories, so this only works for a file under the
    user's own document tree.
    """
    import subprocess
    script = f'''tell application "Microsoft Word"
        open POSIX file "{path}"
        set d to active document
        set pc to 0
        repeat with i from 1 to 40
            try
                set pc to (compute statistics d statistic statistic pages)
                if pc > 0 then exit repeat
            end try
            delay 0.25
        end repeat
        close d saving no
        return pc
    end tell'''
    try:
        out = subprocess.run(["osascript", "-e", script], capture_output=True,
                             text=True, timeout=180)
        return int(out.stdout.strip()) or None
    except (subprocess.TimeoutExpired, ValueError, OSError):
        return None

PT_TO_MM = 25.4 / 72


def paragraph_height_mm(text: str, size_pt: float, leading: float,
                        before_pt: float, after_pt: float, font: ImageFont.FreeTypeFont) -> float:
    """Height one paragraph occupies, including the spacing above and below it.

    Args:
        text: The paragraph's rendered text.
        size_pt: Font size in points.
        leading: Line spacing as a multiple, e.g. 1.40 for 140 per cent.
        before_pt: Space above the paragraph, in points.
        after_pt: Space below, in points.
        font: A font loaded at MEASURE_PT, used only for advance widths.

    Returns:
        Height in millimetres.
    """
    width_pt = font.getlength(text) * (size_pt / MEASURE_PT) * CHAR_SCALE
    width_pt += CHAR_SPACING * size_pt * len(text)
    width_mm = width_pt * PT_TO_MM
    lines = max(1, -(-width_mm // LINE_WIDTH_MM))     # ceiling division
    line_mm = size_pt * LINE_BOX_EM * leading * PT_TO_MM
    return lines * line_mm + (before_pt + after_pt) * PT_TO_MM


def main() -> None:
    if not Path(PROXY_FONT).exists():
        sys.exit(f"proxy font not found: {PROXY_FONT}")
    if not DOCX.exists():
        sys.exit(f"manuscript not found: {DOCX}")

    font = ImageFont.truetype(PROXY_FONT, MEASURE_PT)
    doc = Document(str(DOCX))

    text_mm = 0.0
    for para in doc.paragraphs:
        body = para.text.strip()
        if not body:
            continue
        size = para.runs[0].font.size.pt if para.runs and para.runs[0].font.size else 10.0
        # The abstract is the only style set to 120 per cent leading; everything else is 140.
        leading = 1.20 if size == 10 and len(body) > 800 else 1.40
        pf = para.paragraph_format
        text_mm += paragraph_height_mm(
            body, size, leading,
            pf.space_before.pt if pf.space_before else 0.0,
            pf.space_after.pt if pf.space_after else 0.0,
            font,
        )

    rows = sum(len(t.rows) for t in doc.tables)
    tables_mm = rows * TABLE_ROW_MM + len(doc.tables) * TABLE_PAD_MM
    total_mm = (text_mm + tables_mm) * CALIBRATION
    pages = total_mm / USABLE_HEIGHT_MM

    print(f"text and spacing   {text_mm:8,.0f} mm")
    print(f"tables             {tables_mm:8,.0f} mm   ({len(doc.tables)} tables, {rows} rows)")
    print(f"total              {total_mm:8,.0f} mm   over {USABLE_HEIGHT_MM:.0f} mm per page")
    print(f"                   (includes calibration factor {CALIBRATION})")
    print(f"\nestimate           {pages:.2f} pages   limit is 5")
    if pages > 5.0:
        over = (pages - 4.9) / pages
        doc = Document(str(DOCX))
        chars = sum(len(p.text) for p in doc.paragraphs)
        print(f"OVER by {pages - 5.0:.2f} pages. To reach 4.9, cut about {over:.0%} "
              f"of the content, roughly {int(chars * over):,} characters.")
    real = word_pages(DOCX)
    if real is None:
        print("\nWord could not be queried; the number above is the fallback model and is not "
              "reliable to better than a page.")
    else:
        print(f"\nWord reports {real} pages. That is the number that counts; the estimate above "
              "is only the fallback model.")


if __name__ == "__main__":
    main()
